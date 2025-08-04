import streamlit as st
import json
from docx import Document
from docx.shared import Inches
from pathlib import Path

st.set_page_config(page_title="Module 10: Final Report Generator", layout="wide")
st.title("📄 Module 10: Final Report Generator")

json_path = Path("module9_analysis.json")
if not json_path.exists():
    st.error("module9_analysis.json not found. Please complete Module 9 first.")
    st.stop()

with open(json_path, "r") as f:
    data = json.load(f)

doc = Document()
doc.add_heading("Ichiban Market Valuation Report", 0)

sp = data["subject_property"]
doc.add_heading("Subject Property", level=1)
doc.add_paragraph(f"Address: {sp['address']}")
doc.add_paragraph(f"Above Grade SF: {sp['above_grade_sf']} | Bedrooms: {sp['bedrooms']} | Bathrooms: {sp['bathrooms']}")

doc.add_heading("Valuation Summary", level=1)
doc.add_paragraph(f"Online Estimate Average: ${data['online_estimate_average']:,}")
doc.add_paragraph(f"Adjusted Comp Range: {data['adjusted_price_range'][0]:,} – {data['adjusted_price_range'][1]:,}")
doc.add_paragraph(f"Average PPSF: ${data['average_ppsf']:,}")
days_list = [c.get("days_in_mls") for c in data["comps"] if isinstance(c.get("days_in_mls"), (int, float))]
if days_list:
    doc.add_paragraph(f"Average Days in MLS: {round(sum(days_list)/len(days_list), 1)}")

doc.add_heading("GPT Market Commentary", level=1)
doc.add_paragraph(data.get("gpt_review", "No commentary available."))

doc.add_heading("Adjusted Comparable Sales", level=1)
table = doc.add_table(rows=1, cols=8)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Address"
hdr_cells[1].text = "AG SF"
hdr_cells[2].text = "Net Price"
hdr_cells[3].text = "AG Adj"
hdr_cells[4].text = "BGF Adj"
hdr_cells[5].text = "BGU Adj"
hdr_cells[6].text = "Total Adj"
hdr_cells[7].text = "Adjusted Price"

for c in data["comps"]:
    row = table.add_row().cells
    row[0].text = c["full_address"]
    row[1].text = str(c["ag_sf"])
    row[2].text = f"${c['net_price']:,}"
    row[3].text = f"${c['ag_adj']:,}"
    row[4].text = f"${c.get('bgf_adj', 0):,}"
    row[5].text = f"${c.get('bgu_adj', 0):,}"
    row[6].text = f"${c['total_adjustments']:,}"
    row[7].text = f"${c['adjusted_price']:,}"

output_path = "Ichiban_Valuation_Report.docx"
doc.save(output_path)
with open(output_path, "rb") as f:
    st.download_button("📥 Download Final Report", f, file_name=output_path)

st.success("✅ Final Report Generated")
