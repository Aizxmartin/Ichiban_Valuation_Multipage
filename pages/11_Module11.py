
import json
from pathlib import Path
import statistics as stats

import streamlit as st
from docx import Document

# OpenAI SDK v1+
try:
    from openai import OpenAI
except Exception as e:
    OpenAI = None

st.set_page_config(page_title="Module 11: GPT Market Review (with Auto-Fallback)", page_icon="🤖")
st.title("🤖 Module 11: GPT Commentary & Enhanced Report — with Auto-Fallback")

# --- Load API key (optional; we can still run fallback without it)
api_key = st.secrets.get("general", {}).get("openai_key", "")

# --- Load JSON from Module 8
json_path = Path("module9_analysis.json")
if not json_path.exists():
    st.error("❌ Could not find 'module9_analysis.json'. Please run Module 8 first.")
    st.stop()

with open(json_path, "r") as f:
    summary = json.load(f)

subject = summary.get("subject_property", {})
online_avg = summary.get("online_estimate_average", 0)
raw_low, raw_high = summary.get("adjusted_price_range", [0, 0])
norm_low, norm_high = summary.get("normalized_range", [raw_low, raw_high])
norm_median = summary.get("normalized_median", (norm_low + norm_high) / 2)
avg_ppsf = summary.get("average_ppsf", 0)
days_mls = summary.get("average_days_in_mls", "N/A")
comps = summary.get("comps", [])
outlier_notes = summary.get("outlier_notes", {})
estimates = summary.get("online_estimates", {})

zillow_val = estimates.get("zillow")
redfin_val = estimates.get("redfin")
real_avm_val = estimates.get("real_avm")

# --- Prepare remarks (limit to 5)
remarks_snippets = [c.get("remarks_snippet", "") for c in comps if c.get("remarks_snippet")]
remarks_snippets = remarks_snippets[:5]
remarks_block = "\n".join([f"- {s}" for s in remarks_snippets]) if remarks_snippets else "No notable public remarks provided."

# --- Recommend range inside normalized band (target ≤ 75k)
def propose_list_range(low, high, median):
    target = 75_000.0
    spread = float(high - low)
    if spread <= target:
        return float(low), float(high), False
    half = target / 2.0
    return max(float(low), float(median) - half), min(float(high), float(median) + half), True

rec_low, rec_high, wider_flag = propose_list_range(norm_low, norm_high, norm_median)

# --- Helper: friendly money
def money(v):
    try:
        return f"${float(v):,.0f}"
    except Exception:
        return str(v)

# --- Optional GPT path
client = None
gpt_commentary = None
gpt_error = None

def show_friendly_api_error(e: Exception):
    status = getattr(e, "status_code", None)
    msg = getattr(e, "message", None) or str(e)
    if status == 401:
        st.error("❌ GPT auth failed (401). Falling back to non-GPT summary.")
    elif status == 429:
        st.error("❌ GPT quota/limits (429). Falling back to non-GPT summary.")
    elif status == 403:
        st.error("❌ GPT forbidden (403). Falling back to non-GPT summary.")
    elif status:
        st.error(f"❌ GPT error {status}. Falling back to non-GPT summary.")
    else:
        st.error(f"❌ GPT error. Falling back to non-GPT summary. Details: {msg}")

# --- Build prompt (for when GPT is available)
rules = """
Final Review Rules:
1) Show Raw vs Normalized ranges.
2) Recommend a list-price range INSIDE normalized; aim ≤ $75K spread.
3) If >$75K (up to $150K), justify clearly based on data; avoid >$150K.
4) Use Average Days in MLS to frame pacing.
5) Use comp remarks for qualitative nuance.
"""

prompt = f"""
You are a senior real estate pricing analyst. Produce:
- 4–7 bullets (market position, comp alignment, demand cues, risks)
- Short paragraph with strategy and recommended range

Subject Property:
- Address: {subject.get("address", "N/A")}
- Above Grade SF: {subject.get("ag_sf", "N/A")}
- Bedrooms: {subject.get("bedrooms", "N/A")}
- Bathrooms: {subject.get("bathrooms", "N/A")}
- Below Grade Finished: {subject.get("below_grade_finished", 0)}
- Below Grade Unfinished: {subject.get("below_grade_unfinished", 0)}

Valuation:
- Online Estimate Average: {money(online_avg)}
- Raw Adjusted Comp Range: {money(raw_low)} – {money(raw_high)}
- Normalized Range: {money(norm_low)} – {money(norm_high)} (median {money(norm_median)})
- Average Price Per SF: ${avg_ppsf}
- Average Days in MLS: {days_mls}

Online Estimates:
- Zillow: {money(zillow_val) if zillow_val else "N/A"}
- Redfin: {money(redfin_val) if redfin_val else "N/A"}
- Real AVM: {money(real_avm_val) if real_avm_val else "N/A"}

Comp Remarks (snippets):
{remarks_block}

Outlier Notes:
{json.dumps(outlier_notes, indent=2)}

{rules}

Your response must:
- Recommend a range inside the normalized band: {money(rec_low)} – {money(rec_high)}
- If this still exceeds $75K, explain why it is acceptable or propose a tightened alternative.
""".strip()

st.subheader("🔎 Prompt Preview")
st.code(prompt)

# --- Expose Module 9 backup if it exists
backup_path = Path("Ichiban_Backup_Report_Module9.docx")
if backup_path.exists():
    st.info("📎 A backup non-GPT report from Module 9 is available below.")
    with open(backup_path, "rb") as f:
        st.download_button("📥 Download Backup Report (from Module 9)", f, file_name=backup_path.name)

# --- Run button
if st.button("Run Analysis and Generate Enhanced Report"):
    used_fallback = False
    if api_key and OpenAI is not None:
        try:
            client = OpenAI(api_key=api_key)
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a senior real estate pricing analyst."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.6,
                max_tokens=1000
            )
            gpt_commentary = resp.choices[0].message.content.strip()
        except Exception as e:
            gpt_error = e
            show_friendly_api_error(e)
            used_fallback = True
    else:
        used_fallback = True
        if not api_key:
            st.warning("No API key found. Using fallback deterministic summary.")
        elif OpenAI is None:
            st.warning("OpenAI SDK not installed. Using fallback deterministic summary.")

    # --- Build commentary (fallback if needed)
    if used_fallback or not gpt_commentary:
        # Deterministic, data-driven bullets and paragraph
        bullets = []
        bullets.append(f"Subject sits between the online estimate {money(online_avg)} and the normalized comp band {money(norm_low)}–{money(norm_high)} (median {money(norm_median)}).")
        bullets.append(f"Average PPSF is ${avg_ppsf}, providing a broad sanity check across adjusted comps.")
        dm = str(days_mls)
        if dm != "N/A":
            bullets.append(f"Average Days in MLS ≈ {dm}; pricing should reflect desired time-to-contract relative to market pace.")
        else:
            bullets.append("Average Days in MLS unavailable; recommend conservative pricing to avoid extended market time.")
        if remarks_snippets:
            bullets.append("Comp remarks mention notable features/upgrades; consider their presence/absence for the subject when positioning.")
        spread = rec_high - rec_low
        if spread > 75000:
            bullets.append("Recommended range initially exceeds the $75K target; propose justification or tighten around the median as needed.")
        else:
            bullets.append("Recommended range meets the ≤ $75K target for clear buyer messaging.")
        para = (
            f"Given the filtered, outlier-aware range of {money(norm_low)}–{money(norm_high)} and the online estimate at {money(online_avg)}, "
            f"a recommended listing window of {money(rec_low)}–{money(rec_high)} balances buyer reach with price protection. "
            f"If market feedback is slow relative to the average days on market, adjust downward toward the lower half of the band; "
            f"if activity is brisk, test near the upper half."
        )
        gpt_commentary = "\n".join([f"- {b}" for b in bullets]) + "\n\n" + para

    # --- Build DOCX
    doc = Document()
    title = "Enhanced Market Valuation Summary (Fallback)" if used_fallback else "Enhanced Market Valuation Summary"
    doc.add_heading(title, 0)

    # Subject header
    doc.add_heading("Subject Property Summary", level=1)
    doc.add_paragraph(f"Address: {subject.get('address', 'N/A')}")
    doc.add_paragraph(f"Above Grade SF: {subject.get('ag_sf', 'N/A')}")
    doc.add_paragraph(f"Bedrooms: {subject.get('bedrooms', 'N/A')}")
    doc.add_paragraph(f"Bathrooms: {subject.get('bathrooms', 'N/A')}")
    doc.add_paragraph(f"Below Grade Finished: {subject.get('below_grade_finished', 0)}")
    doc.add_paragraph(f"Below Grade Unfinished: {subject.get('below_grade_unfinished', 0)}")

    # Overview
    doc.add_heading("Enhanced Property Overview", level=1)
    if used_fallback and gpt_error:
        doc.add_paragraph("(GPT unavailable — auto-generated deterministic summary)")
    doc.add_paragraph(gpt_commentary)

    # Online estimates
    doc.add_heading("Online Estimates", level=1)
    est = doc.add_table(rows=4, cols=2)
    est.cell(0, 0).text = "Source"
    est.cell(0, 1).text = "Value"
    est.cell(1, 0).text = "Zillow"
    est.cell(1, 1).text = money(zillow_val) if zillow_val else "N/A"
    est.cell(2, 0).text = "Redfin"
    est.cell(2, 1).text = money(redfin_val) if redfin_val else "N/A"
    est.cell(3, 0).text = "Real AVM"
    est.cell(3, 1).text = money(real_avm_val) if real_avm_val else "N/A"

    # Adjusted comps table with Total Adjustments
    doc.add_heading("Adjusted Comparable Properties", level=1)
    comp_table = doc.add_table(rows=1, cols=4)
    hdrs = comp_table.rows[0].cells
    hdrs[0].text = "Address"
    hdrs[1].text = "AG SF"
    hdrs[2].text = "Total Adjustments"
    hdrs[3].text = "Adjusted Price"
    for c in comps:
        row = comp_table.add_row().cells
        row[0].text = str(c.get("full_address", ""))
        row[1].text = str(c.get("ag_sf", ""))
        ta = c.get("total_adjustments", (c.get("ag_adj", 0) + c.get("bgf_adj", 0) + c.get("bgu_adj", 0)))
        row[2].text = money(ta)
        row[3].text = money(c.get("adjusted_price", 0))

    # Ranges & Recommendation
    doc.add_heading("Ranges", level=1)
    doc.add_paragraph(f"Raw Adjusted Comp Range: {money(raw_low)} – {money(raw_high)}")
    doc.add_paragraph(f"Normalized Range (outlier-aware): {money(norm_low)} – {money(norm_high)} (median {money(norm_median)})")
    doc.add_heading("Recommended Market Range", level=1)
    doc.add_paragraph(f"Recommended Listing Range: {money(rec_low)} – {money(rec_high)}")
    doc.add_paragraph("Policy: Target ≤ $75K spread. Wider ranges require explicit data-based rationale; avoid > $150K unless compelling reasons exist.")

    # Save & serve
    out_name = "Enhanced_GPT_Market_Report_AutoFallback.docx"
    with open(out_name, "wb") as f:
        doc.save(f)

    with open(out_name, "rb") as f:
        st.download_button("📥 Download Enhanced Report (with Auto-Fallback)", f, file_name=out_name)

    # Also expose the Module 9 backup again (if available)
    if backup_path.exists():
        with open(backup_path, "rb") as f:
            st.download_button("📎 Download Backup Report (Module 9)", f, file_name=backup_path.name)
