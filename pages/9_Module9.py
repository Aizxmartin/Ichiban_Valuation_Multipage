
import json
from pathlib import Path
from docx import Document
import streamlit as st

st.set_page_config(page_title="Module 9: Backup Report (Schema-Aware + UX)", layout="wide")
st.title("🧰 Module 9: Backup Non-GPT Report — Schema Aligned (No API Required)")

json_path = Path("module9_analysis.json")
if not json_path.exists():
    st.error("❌ Could not find 'module9_analysis.json'. Please run Module 8 first.")
    st.stop()

with open(json_path, "r") as f:
    data = json.load(f)

# Extract with safe defaults
sp = data.get("subject_property", {})
online_avg = data.get("online_estimate_average", 0)
raw_low, raw_high = data.get("adjusted_price_range", [0, 0])
norm_low, norm_high = data.get("normalized_range", [raw_low, raw_high])
norm_median = data.get("normalized_median", (norm_low + norm_high)/2)
avg_ppsf = data.get("average_ppsf", 0)
avg_days = data.get("average_days_in_mls", "N/A")
online_estimates = data.get("online_estimates", {})
comps = data.get("comps", [])

def money(v):
    try:
        return f"${float(v):,.0f}"
    except Exception:
        return str(v)

# UX guidance
st.info("🧷 This is your no-API backup path. If Module 11/GPT fails, this report will still be complete and consistent with Module 8.")

# Quick validations
if not online_estimates:
    st.warning("ℹ️ Online Estimates are missing in JSON. Add them in Module 3 so they appear here and in GPT reports.")
if avg_days == "N/A":
    st.warning("ℹ️ Average Days in MLS is N/A in JSON. Add a 'Days In MLS' column in your comps CSV for better commentary.")
if len(comps) == 0:
    st.error("No comps present in JSON. Ensure Module 8 ran successfully with adjusted comps.")
    st.stop()

# Build DOCX
doc = Document()
doc.add_heading("Ichiban Market Valuation Report — Backup (Non-GPT)", 0)

# Subject Property
doc.add_heading("Subject Property", level=1)
doc.add_paragraph(f"Address: {sp.get('address', 'N/A')}")
doc.add_paragraph(f"Above Grade SF: {sp.get('ag_sf', 'N/A')} | Bedrooms: {sp.get('bedrooms', 'N/A')} | Bathrooms: {sp.get('bathrooms', 'N/A')}")
doc.add_paragraph(f"Below Grade Finished: {sp.get('below_grade_finished', 'N/A')} | Below Grade Unfinished: {sp.get('below_grade_unfinished', 'N/A')}")

# Valuation Summary
doc.add_heading("Valuation Summary", level=1)
doc.add_paragraph(f"Online Estimate Average: {money(online_avg)}")
doc.add_paragraph(f"Average Price per SF: ${avg_ppsf}")
doc.add_paragraph(f"Average Days in MLS: {avg_days}")

# Ranges
doc.add_heading("Pricing Ranges", level=1)
doc.add_paragraph(f"Raw Adjusted Comp Range: {money(raw_low)} – {money(raw_high)}")
doc.add_paragraph(f"Normalized Range (outlier-aware): {money(norm_low)} – {money(norm_high)} (median {money(norm_median)})")
if abs(raw_low - norm_low) < 1 and abs(raw_high - norm_high) < 1:
    doc.add_paragraph("Note: Normalized range equals raw range; no strong outliers were detected by the filters.")

# Online Estimates
doc.add_heading("Online Estimates (Manual Inputs)", level=1)
est_table = doc.add_table(rows=4, cols=2)
est_table.cell(0, 0).text = "Source"
est_table.cell(0, 1).text = "Value"
z = online_estimates.get("zillow")
r = online_estimates.get("redfin")
a = online_estimates.get("real_avm")
est_table.cell(1, 0).text = "Zillow"; est_table.cell(1, 1).text = money(z) if z is not None else "N/A"
est_table.cell(2, 0).text = "Redfin"; est_table.cell(2, 1).text = money(r) if r is not None else "N/A"
est_table.cell(3, 0).text = "Real AVM"; est_table.cell(3, 1).text = money(a) if a is not None else "N/A"

# Adjusted Comps
doc.add_heading("Adjusted Comparable Sales (with Total Adjustments)", level=1)
table_cols = ["Address", "AG SF", "Net Price", "AG Adj", "BGF Adj", "BGU Adj", "Total Adj", "Adjusted Price", "Days MLS"]
table = doc.add_table(rows=1, cols=len(table_cols))
for i, h in enumerate(table_cols):
    table.rows[0].cells[i].text = h

for c in comps:
    row = table.add_row().cells
    row[0].text = str(c.get("full_address", ""))
    row[1].text = str(c.get("ag_sf", ""))
    row[2].text = money(c.get("net_price"))
    row[3].text = money(c.get("ag_adj", 0))
    row[4].text = money(c.get("bgf_adj", 0))
    row[5].text = money(c.get("bgu_adj", 0))
    total_adj = c.get("total_adjustments", (c.get("ag_adj", 0) + c.get("bgf_adj", 0) + c.get("bgu_adj", 0)))
    row[6].text = money(total_adj)
    row[7].text = money(c.get("adjusted_price"))
    dim = c.get("days_in_mls", "N/A")
    row[8].text = "N/A" if dim is None else str(dim)

# Save + Download
output_path = "Ichiban_Backup_Report_Module9.docx"
doc.save(output_path)
with open(output_path, "rb") as f:
    st.download_button("📥 Download Backup Report (Module 9)", f, file_name=output_path)

st.success("✅ Backup Report Created (Module 9)")
st.caption("Tip: For the GPT-enhanced narrative, open **Module 11**. It will auto-fallback if the API is unavailable.")
